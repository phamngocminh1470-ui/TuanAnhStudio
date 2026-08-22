import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def style_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.79)     # ~2.0 cm
        section.bottom_margin = Inches(0.79)  # ~2.0 cm
        section.left_margin = Inches(1.18)    # ~3.0 cm
        section.right_margin = Inches(0.79)   # ~2.0 cm

    # Style default Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)

def convert_markdown_to_docx(md_path, docx_path, doc_title):
    doc = Document()
    style_document(doc)

    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        
        # Determine cols
        num_cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, "B0C4DE", "6", "single")

        for r_idx, row_data in enumerate(table_rows):
            is_header = (r_idx == 0)
            for c_idx in range(num_cols):
                cell = table.cell(r_idx, c_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
                
                text_val = row_data[c_idx] if c_idx < len(row_data) else ""
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                
                if is_header:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_background(cell, "EBF3FB")
                    run = p.add_run(text_val)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
                else:
                    run = p.add_run(text_val)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if c_idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    if r_idx % 2 == 1:
                        set_cell_background(cell, "FAFCFF")

        doc.add_paragraph() # Spacing after table
        table_rows = []
        in_table = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        stripped = line.strip()

        # Check table
        if stripped.startswith('|') and stripped.endswith('|'):
            # Table row
            if '---' in stripped and ('|' in stripped):
                # Separator line, ignore
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            # Clean formatting in cells
            clean_cells = []
            for c in cells:
                c_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', c)
                c_clean = re.sub(r'\*(.*?)\*', r'\1', c_clean)
                c_clean = re.sub(r'\$(.*?)\$', r'\1', c_clean)
                clean_cells.append(c_clean)
            table_rows.append(clean_cells)
            in_table = True
            i += 1
            continue
        elif in_table:
            flush_table()

        if not stripped:
            i += 1
            continue

        # Headers
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped[2:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[3:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81) # Deep Blue
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[4:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x22, 0x33, 0x44)
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[5:].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.italic = True
        elif stripped.startswith('> '):
            # Quote / Callout
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            text_content = stripped[2:].strip()
            
            # Format bold and italic
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text_content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12.5)
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.font.italic = True
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12.5)
                else:
                    r = p.add_run(part)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12.5)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif stripped.startswith('```'):
            # Code block or diagram
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\r\n'))
                i += 1
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
            code_text = "\n".join(code_lines)
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x1E, 0x39, 0x5B)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.3
            text_content = stripped[2:].strip()
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$)', text_content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.font.italic = True
                elif part.startswith('$') and part.endswith('$'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Cambria Math'
                    r.font.italic = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
        elif re.match(r'^\d+\.\s', stripped):
            # Numbered list
            match = re.match(r'^(\d+\.\s)(.*)$', stripped)
            prefix = match.group(1)
            text_content = match.group(2)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.3
            
            p.add_run(prefix).font.bold = True
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$)', text_content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.font.italic = True
                elif part.startswith('$') and part.endswith('$'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Cambria Math'
                    r.font.italic = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
        elif stripped == '---':
            # Separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('________________________________________________________')
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(9)
        else:
            # Normal paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.3
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.font.italic = True
                elif part.startswith('$') and part.endswith('$'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Cambria Math'
                    r.font.italic = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)

        i += 1

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")

if __name__ == '__main__':
    files_to_convert = [
        ("HO_SO_KHKT_01_BAO_CAO_TOAN_VAN.md", "HO_SO_KHKT_01_BAO_CAO_TOAN_VAN.docx", "BÁO CÁO TOÀN VĂN KHKT"),
        ("HO_SO_KHKT_02_TOM_TAT_DU_AN.md", "HO_SO_KHKT_02_TOM_TAT_DU_AN.docx", "TÓM TẮT DỰ ÁN KHKT"),
        ("HO_SO_KHKT_04_KICH_BAN_THUYT_TRINH.md", "HO_SO_KHKT_04_KICH_BAN_THUYT_TRINH.docx", "KỊCH BẢN THUYẾT TRÌNH KHKT"),
        ("HO_SO_KHKT_05_BO_CAU_HOI_PHAN_BIEN_GIAM_KHAO.md", "HO_SO_KHKT_05_BO_CAU_HOI_PHAN_BIEN_GIAM_KHAO.docx", "BỘ CÂU HỎI PHẢN BIỆN GIÁM KHẢO"),
    ]

    for md_file, docx_file, title in files_to_convert:
        convert_markdown_to_docx(md_file, docx_file, title)

    print("ALL KHKT WORD DOCUMENTS GENERATED SUCCESSFULLY!")
