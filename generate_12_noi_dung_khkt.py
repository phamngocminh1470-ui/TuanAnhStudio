import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

OUTPUT_DIR = "HO_SO_BAO_CAO_KHKT"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def apply_so_page_setup(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.79)     # 2.0 cm
        section.bottom_margin = Inches(0.79)  # 2.0 cm
        section.left_margin = Inches(1.18)    # 3.0 cm
        section.right_margin = Inches(0.79)   # 2.0 cm

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="B0C4DE", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_header(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x0F, 0x36, 0x66)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        run.font.size = Pt(14)
        run.font.italic = True
    return p

def add_paragraph_text(doc, text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(14)
        r_pre.font.bold = True
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.italic = italic
    return p

def add_bullet_point(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(14)
        r_pre.font.bold = True
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    return p

def build_12_core_topics_doc():
    doc = Document()
    apply_so_page_setup(doc)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_title.add_run("TÀI LIỆU ÔN TẬP VÀ BẢO VỆ ĐỀ TÀI KHOA HỌC KỸ THUẬT\n")
    r_sub.font.size = Pt(13)
    r_sub.font.bold = True

    r_main = p_title.add_run("TỔNG HỢP 12 NỘI DUNG TRỌNG TÂM CẦN CHUẨN BỊ\nTRONG HỒ SƠ VÀ BẢO VỆ TRƯỚC HỘI ĐỒNG GIÁM KHẢO\n")
    r_main.font.size = Pt(16)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(0x0F, 0x36, 0x66)

    r_proj = p_title.add_run("Dự án: Hệ thống học tập thích ứng cá nhân hóa hỗ trợ tự học tiếng Anh cho học sinh THPT dựa trên mô hình IRT, thuật toán SM-2 và Trí tuệ nhân tạo\nWebsite: https://tuananhstudio.top")
    r_proj.font.size = Pt(13)
    r_proj.font.italic = True

    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.add_run("________________________________________________________\n").font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # 1. TÍNH MỚI
    add_header(doc, "1. TÍNH MỚI CỦA ĐỀ TÀI (NOVELTY & INNOVATION)", level=1)
    add_bullet_point(doc, " Là công trình đầu tiên tại Việt Nam tích hợp đồng bộ 3 trụ cột khoa học: Toán học đo lường giáo dục hiện đại (Mô hình 3PL IRT), Tâm lý học nhận thức thần kinh (Thuật toán lặp ngắt quãng SuperMemo-2) và Trí tuệ nhân tạo đàm thoại gợi mở (Socratic AI Scaffolding) vào một nền tảng tự học tiếng Anh hoàn chỉnh cho học sinh phổ thông.", "Đột phá về tích hợp liên ngành: ")
    add_bullet_point(doc, " Xóa bỏ cơ chế kiểm tra cào bằng truyền thống. Hệ thống không tính điểm tổng thô mà đo lường năng lực tiềm ẩn Theta của học sinh sau từng câu hỏi, tự động đẩy câu tiếp theo có độ khó tối ưu theo hàm thông tin Fisher tối đa.", "Chấm thi và lộ trình thích ứng: ")
    add_bullet_point(doc, " Module nhận diện âm thanh đối chiếu âm vị IPA, chỉ ra chính xác lỗi nuốt âm đuôi (/s/, /ed/, /θ/, /t/) và cho phép học sinh bấm trực tiếp vào từng từ bị sai để nghe phát âm đọc chậm riêng từ đó để nhại theo và sửa ngay lập tức.", "Tương tác ngữ âm trực quan: ")

    # 2. TÍNH KHOA HỌC
    add_header(doc, "2. TÍNH KHOA HỌC (SCIENTIFIC RIGOR & THEORETICAL BASIS)", level=1)
    add_bullet_point(doc, " Dựa trên lý thuyết trắc nghiệm cổ điển và hiện đại (Lord, 1980) với hàm xác suất logistic 3 tham số P_i(Theta) = c_i + (1 - c_i) / (1 + exp(-1.7 * a_i * (Theta - b_i))) và thuật toán tích phân số Gauss-Hermite 21 điểm nút để tính kỳ vọng năng lực EAP.", "Cơ sở toán học đo lường giáo dục: ")
    add_bullet_point(doc, " Dựa trên quy luật suy giảm trí nhớ của Hermann Ebbinghaus (1885) và thuật toán tối ưu hóa ghi nhớ SuperMemo-2 (P.A. Wozniak, 1990) để tự động cập nhật Hệ số Dễ nhớ (Easiness Factor) và khoảng cách ngày ôn tập I(n).", "Cơ sở tâm lý học nhận thức: ")
    add_bullet_point(doc, " Áp dụng Thuyết vùng phát triển gần nhất (ZPD - Vygotsky) và phương pháp giàn giáo sư phạm Socrates (Scaffolding): AI không bao giờ giải hộ hay đưa sẵn đáp án mà đặt câu hỏi phản biện, dẫn dắt học sinh tự tìm ra cách sửa.", "Cơ sở lý luận dạy học: ")
    add_bullet_point(doc, " Toàn bộ số liệu thực nghiệm được xử lý bằng các phép kiểm định thống kê toán học nghiêm ngặt: Independent Samples t-test, Paired Samples t-test, và Đo lường độ lớn ảnh hưởng chuẩn hóa Cohen's d.", "Phương pháp xử lý số liệu chuẩn xác: ")

    # 3. TÍNH THỰC TIỄN
    add_header(doc, "3. TÍNH THỰC TIỄN (PRACTICAL EFFECTIVENESS)", level=1)
    add_bullet_point(doc, " Giải quyết triệt để 3 rào cản nhức nhối của học sinh phổ thông: (1) Quá tải bài tập không vừa sức; (2) Học vẹt từ vựng rồi quên 80% sau 1-2 tuần; (3) E ngại nuốt âm đuôi và thiếu môi trường tương tác 1-1.", "Giải quyết đúng vấn đề thực tế: ")
    add_bullet_point(doc, " Kết quả thực nghiệm 120 học sinh trong 8 tuần cho thấy nhóm dùng hệ thống tăng trung bình +2.45 điểm (gấp 3.3 lần nhóm đối chứng), nhớ từ vựng sau 14 ngày đạt 84.5% (tăng +105%), tiết kiệm 52.4% thời gian làm bài kiểm tra.", "Hiệu quả thực chứng vượt bậc: ")
    add_bullet_point(doc, " Sản phẩm không nằm trên giấy tờ lý thuyết mà đã chạy thực tế trực tuyến 100% tại https://tuananhstudio.top với độ trễ phản hồi dưới 0.5 giây.", "Sản phẩm hoàn thiện, sẵn sàng vận hành: ")

    # 4. TÍNH CỘNG ĐỒNG
    add_header(doc, "4. TÍNH CỘNG ĐỒNG VÀ NHÂN VĂN (COMMUNITY IMPACT)", level=1)
    add_bullet_point(doc, " Cung cấp giải pháp học tập miễn phí 100% cho mọi học sinh và nhà trường, xóa bỏ rào cản học thêm tốn kém.", "Bình đẳng cơ hội giáo dục: ")
    add_bullet_point(doc, " Hệ thống được tối ưu hóa cực kỳ nhẹ (PageSpeed 99/100, Mobile 89/100), hoạt động mượt mà trên cả điện thoại thông minh bình dân và mạng di động 3G/4G yếu, giúp học sinh vùng sâu, vùng xa đều có thể tiếp cận gia sư AI chất lượng cao.", "Khả năng tiếp cận phổ quát: ")
    add_bullet_point(doc, " Cung cấp Bảng Quản trị Nghiên cứu Sư phạm (Admin Research Panel) cho giáo viên theo dõi tiến độ của học sinh và xuất dữ liệu Excel/CSV phục vụ nghiên cứu sư phạm của nhà trường.", "Hỗ trợ giáo viên và nhà trường: ")

    # 5. LÍ DO CHỌN DỰ ÁN
    add_header(doc, "5. LÍ DO CHỌN DỰ ÁN (RATIONALE & BACKGROUND)", level=1)
    add_paragraph_text(doc, "Chương trình GDPT 2018 và định dạng đề thi Tốt nghiệp THPT mới từ năm 2025 đòi hỏi học sinh phát triển năng lực ngôn ngữ thực chất. Tuy nhiên, trong thực tế trường phổ thông hiện nay:")
    add_bullet_point(doc, " Sĩ số lớp đông (40-45 học sinh/lớp) khiến việc dạy học mang tính đồng loạt cào bằng; học sinh khá tốn thời gian vào bài dễ, học sinh yếu bị ngợp trước bài khó. Giáo viên không thể kèm cặp sát sao từng lỗ hổng ngữ pháp riêng lẻ.", "Rào cản về sĩ số và thời gian: ")
    add_bullet_point(doc, " Học sinh chủ yếu chép danh sách từ để kiểm tra 15 phút rồi quên 70-80% chỉ sau 1-2 tuần do thiếu cơ chế lặp lại ngắt quãng.", "Thói quen học vẹt mau quên: ")
    add_bullet_point(doc, " Đa số học sinh rất e ngại phần phát âm vì hay bị nuốt âm đuôi và thiếu môi trường luyện tập đối thoại 1-1 thường xuyên.", "Tâm lý e ngại phát âm: ")
    add_paragraph_text(doc, "Xuất phát từ những trăn trở thực tiễn đó, nhóm chúng em đã nghiên cứu và phát triển dự án AI English Mentor.")

    # 6. MỤC ĐÍCH NGHIÊN CỨU
    add_header(doc, "6. MỤC ĐÍCH NGHIÊN CỨU (RESEARCH OBJECTIVES)", level=1)
    add_bullet_point(doc, " Xây dựng một nền tảng Web học tập trực tuyến thông minh, thích ứng cá nhân hóa hoàn toàn miễn phí cho học sinh THPT.", "1. Mục tiêu công nghệ: ")
    add_bullet_point(doc, " Ứng dụng mô hình 3PL IRT và thuật toán EAP để ước lượng chính xác năng lực tiềm ẩn Theta và tự động chọn bài tập vừa đúng sức của học sinh.", "2. Mục tiêu đo lường thích ứng: ")
    add_bullet_point(doc, " Ứng dụng thuật toán SuperMemo-2 để chuyển hóa từ vựng ngắn hạn thành trí nhớ dài hạn bền vững.", "3. Mục tiêu củng cố trí nhớ: ")
    add_bullet_point(doc, " Tích hợp phân tích âm học và gia sư Socrates AI để rèn luyện kỹ năng phát âm chuẩn IPA và tư duy tự sửa lỗi ngữ pháp.", "4. Mục tiêu tương tác ngôn ngữ: ")
    add_bullet_point(doc, " Kiểm chứng thực nghiệm tính hiệu quả sư phạm của hệ thống trên 120 học sinh THPT trong 8 tuần.", "5. Mục tiêu thực nghiệm: ")

    # 7. CÂU HỎI NGHIÊN CỨU
    add_header(doc, "7. CÂU HỎI NGHIÊN CỨU (RESEARCH QUESTIONS)", level=1)
    add_paragraph_text(doc, "• Câu hỏi 1 (Q1 - Đo lường): Mô hình trắc nghiệm thích ứng 3PL IRT có giúp xác định chính xác năng lực thực chất và rút ngắn thời gian làm bài của học sinh so với bài thi 50 câu cố định truyền thống không?")
    add_paragraph_text(doc, "• Câu hỏi 2 (Q2 - Trí nhớ): Thuật toán lặp ngắt quãng SuperMemo-2 có nâng cao tỷ lệ duy trì trí nhớ từ vựng sau 14 ngày so với phương pháp học thuộc lòng thông thường không?")
    add_paragraph_text(doc, "• Câu hỏi 3 (Q3 - Phát âm & Tư duy): Module chẩn đoán âm học và Gia sư gợi mở Socrates AI có giúp học sinh tự sửa lỗi nuốt âm đuôi và tăng sự tự tin khi tự học tiếng Anh không?")

    # 8. GIẢ THUYẾT KHOA HỌC
    add_header(doc, "8. GIẢ THUYẾT KHOA HỌC (SCIENTIFIC HYPOTHESES)", level=1)
    add_paragraph_text(doc, "• Giả thuyết H1 (Về điểm số): Học sinh sử dụng hệ thống AI English Mentor (Nhóm Thực nghiệm) sẽ có mức tăng trưởng điểm số kiểm tra và năng lực Theta cao hơn có ý nghĩa thống kê (p < 0.05) so với học sinh học theo phương pháp truyền thống (Nhóm Đối chứng).")
    add_paragraph_text(doc, "• Giả thuyết H2 (Về trí nhớ từ vựng): Tỷ lệ ghi nhớ từ vựng sau 14 ngày của nhóm sử dụng Flashcards SM-2 đạt trên 75%, cao hơn ít nhất 30% so với nhóm đối chứng học truyền thống.")
    add_paragraph_text(doc, "• Giả thuyết H3 (Về thời gian đánh giá): Bài kiểm tra thích ứng CAT dựa trên IRT giảm ít nhất 40% thời gian làm bài so với bài thi 50 câu cố định mà vẫn duy trì sai số ước lượng chuẩn SEM < 0.30.")

    # 9. VẤN ĐỀ NGHIÊN CỨU
    add_header(doc, "9. VẤN ĐỀ NGHIÊN CỨU (CORE RESEARCH PROBLEMS)", level=1)
    add_bullet_point(doc, " Làm thế nào để ước lượng chính xác năng lực thực chất Theta của học sinh ngay sau từng câu hỏi mà không gây độ trễ giao diện?", "Vấn đề 1 (Toán học ước lượng): ")
    add_bullet_point(doc, " Làm thế nào để tự động tính toán chu kỳ lặp lại ngắt quãng phù hợp với tốc độ ghi nhớ riêng biệt của từng cá nhân?", "Vấn đề 2 (Tâm lý học nhận thức): ")
    add_bullet_point(doc, " Làm thế nào để công nghệ AI đóng vai trò người thầy gợi mở tư duy (scaffolding) mà không đưa sẵn đáp án cho học sinh ỷ lại?", "Vấn đề 3 (Sư phạm đàm thoại): ")
    add_bullet_point(doc, " Làm thế nào để chẩn đoán chính xác lỗi âm vị (nhất là âm đuôi tiếng Anh) trên thiết bị di động bình dân trong điều kiện môi trường có tạp âm?", "Vấn đề 4 (Xử lý âm học): ")

    # 10. PHƯƠNG PHÁP NGHIÊN CỨU
    add_header(doc, "10. PHƯƠNG PHÁP NGHIÊN CỨU VÀ CÔNG NGHỆ (METHODOLOGY & TECH)", level=1)
    add_bullet_point(doc, " Nghiên cứu các công trình đo lường giáo dục của Lord (1980), đường cong quên lãng Ebbinghaus (1885), thuật toán SM-2 của Wozniak (1990) và khung năng lực CEFR B1/B2.", "1. Phương pháp nghiên cứu lý thuyết: ")
    add_bullet_point(doc, " Xây dựng hàm 3PL IRT; thuật toán tích phân số Gauss-Hermite 21 điểm nút; thuật toán lặp ngắt quãng SM-2; pipeline phân tích âm học đa tầng; kỹ thuật Structured Few-shot Prompting với JSON schema validation.", "2. Phương pháp mô hình hóa toán học & lập trình: ")
    add_bullet_point(doc, " Thiết kế thực nghiệm có nhóm đối chứng (Pre-test / Post-test Control Group Design) trên 120 học sinh THPT trong 8 tuần liên tục.", "3. Phương pháp thực nghiệm sư phạm: ")
    add_bullet_point(doc, " Sử dụng kiểm định t-test độc lập (Independent t-test), kiểm định t-test theo cặp (Paired t-test), đo lường độ lớn ảnh hưởng Cohen's d và xử lý bằng Python Scipy Stats / SPSS.", "4. Phương pháp thống kê toán học: ")

    # 11. THỰC NGHIỆM SƯ PHẠM
    add_header(doc, "11. THỰC NGHIỆM SƯ PHẠM VÀ KẾT QUẢ ĐỊNH LƯỢNG (EXPERIMENT & DATA)", level=1)
    add_paragraph_text(doc, "Thực nghiệm được tiến hành trong 8 tuần (từ 05/01/2026 đến 28/02/2026) trên 120 học sinh THPT chia đều thành 2 nhóm:")
    
    # Table Results
    table_res = doc.add_table(rows=7, cols=5)
    table_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_res)

    headers = ["Chỉ số đánh giá", "Nhóm Đối chứng (N=60)", "Nhóm Thực nghiệm (N=60)", "Chênh lệch (Delta)", "Ý nghĩa thống kê"]
    for idx, h in enumerate(headers):
        cell = table_res.cell(0, idx)
        cell.text = h
        set_cell_background(cell, "EAF2F8")
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data_rows = [
        ["Điểm Pre-test (Trước TN)", "5.38 ± 1.12", "5.41 ± 1.08", "+0.03", "t = 0.15, p = 0.881 > 0.05 (Tương đồng)"],
        ["Điểm Post-test (Sau TN)", "6.12 ± 1.05", "7.86 ± 0.89", "+1.74", "t = 9.78, p < 0.0001 (Bác bỏ H0)"],
        ["Mức tăng trung bình", "+0.74 điểm", "+2.45 điểm", "Gấp 3.3 lần", "Nhóm TN bứt phá vượt bậc"],
        ["Năng lực Theta đầu ra", "+0.18 ± 0.45", "+0.92 ± 0.38", "+0.74", "p < 0.0001 (Hội tụ chuẩn)"],
        ["Từ vựng nhớ sau 14 ngày", "41.2% (41/100)", "84.5% (84/100)", "+43.3%", "Tăng gấp đôi (+105%)"],
        ["Thời gian làm bài CAT", "45.0 phút (cố định)", "21.4 ± 3.2 phút", "Giảm 52.4%", "Tiết kiệm 23.6 phút"]
    ]

    for r_idx, row in enumerate(data_rows):
        for c_idx, val in enumerate(row):
            cell = table_res.cell(r_idx + 1, c_idx)
            cell.text = val
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(12)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_background(cell, "F9FBFC")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_paragraph_text(doc, "• Kết luận thống kê: Kiểm định t-test độc lập cho kết quả t(118) = 9.78, p < 0.0001 và Cohen's d = 1.79 (> 0.80 - mức độ ảnh hưởng cực lớn), chứng minh giả thuyết H1, H2, H3 là hoàn toàn đúng đắn.")

    # 12. KHẢO SÁT THỰC TẾ
    add_header(doc, "12. KHẢO SÁT THỰC TẾ VÀ ĐÁNH GIÁ MỨC ĐỘ HÀI LÒNG (SURVEY & ACCEPTANCE)", level=1)
    add_paragraph_text(doc, "Khảo sát định lượng trên thang đo Likert 5 mức độ (từ 1: Rất không đồng ý đến 5: Hoàn toàn đồng ý) đối với 60 học sinh Nhóm Thực nghiệm sau 8 tuần:")
    add_bullet_point(doc, " 98.3% học sinh đồng ý hệ thống giúp nâng cao sự tự tin và hứng thú khi tự học môn Tiếng Anh.", "1. Về sự tự tin & động lực: ")
    add_bullet_point(doc, " 100% học sinh khẳng định tính năng Flashcards SM-2 giúp ghi nhớ từ vựng lâu hơn nhiều so với việc chép danh sách từ truyền thống.", "2. Về hiệu quả ghi nhớ từ vựng: ")
    add_bullet_point(doc, " 95.0% học sinh đánh giá cao tính năng bấm trực tiếp vào từ phát âm sai để nghe đọc chậm từng âm và sửa lại tức thì.", "3. Về tính năng sửa phát âm: ")
    add_bullet_point(doc, " 96.7% học sinh thích thú với phương pháp gia sư Socrates AI gợi mở câu hỏi thay vì đưa sẵn đáp án như các công cụ giải bài thông thường.", "4. Về phong cách sư phạm AI: ")
    add_bullet_point(doc, " 98.3% học sinh hài lòng với tốc độ tải trang tức thì và giao diện mượt mà trên điện thoại thông minh.", "5. Về trải nghiệm công nghệ: ")

    out_docx = os.path.join(OUTPUT_DIR, "07_DE_CUONG_CHUAN_BI_12_NOI_DUNG_TRONG_TAM_KHKT.docx")
    doc.save(out_docx)
    print(f"[OK] Generated Word: {out_docx}")

if __name__ == '__main__':
    build_12_core_topics_doc()
